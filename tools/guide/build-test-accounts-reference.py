from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "output" / "test-environment" / "BBC_SMS_TEST_ACCOUNTS_REFERENCE_2026-08-31.docx"
QA_PASSWORD = os.environ.get("QA_PASSWORD", "Fourni séparément par l’administrateur")

NAVY = "173552"
TEAL = "138A83"
GOLD = "D8A72D"
INK = "1C2D42"
MUTED = "64748B"
PALE_TEAL = "E9F7F5"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF5D9"
WHITE = "FFFFFF"
GRID = "C9D5E3"


GROUPS = [
    (
        "Administration et opérations",
        [
            (
                "Administrateur\nadmin",
                "Accès complet à toute l’école : élèves, personnel, pédagogie, finances, paramètres et contrôle des accès.",
                "Vérifier les configurations globales et comparer les droits des autres comptes.",
            ),
            (
                "Comptable global\nqa.accountant.global",
                "Tous les parcours. Peut inscrire des élèves et gérer toute la finance. Personnel en lecture seule. Aucun accès académique, présence, emploi du temps, paramètres ou contrôle des accès.",
                "Tester un paiement, la trésorerie, le compte élève, la paie et l’interdiction d’ouvrir Académique.",
            ),
            (
                "Vie scolaire / préfet\nqa.prefect.walkthrough",
                "Lecture globale des élèves, présence, discipline, alertes, correspondance, aperçu académique et emploi du temps en lecture seule. Pas d’inscription, finance, personnel ou paramètres.",
                "Tester l’appel, un incident disciplinaire et confirmer l’absence du bouton Nouvel élève.",
            ),
        ],
    ),
    (
        "Direction : comparer les périmètres du principal",
        [
            (
                "Principal maternelle\nqa.principal.maternelle",
                "Uniquement la maternelle : élèves, classes, personnel, académique, présence et emploi du temps de ce parcours.",
                "Confirmer qu’aucune classe primaire ou secondaire n’apparaît.",
            ),
            (
                "Principal primaire\nqa.principal.primary",
                "Uniquement le primaire. Peut inscrire des élèves et gérer les opérations du primaire. Finance et trésorerie restent en lecture seule.",
                "Créer un élève primaire puis essayer d’ouvrir une classe secondaire.",
            ),
            (
                "Principal secondaire\nqa.principal.secondary",
                "Tout le secondaire, toutes sections autorisées. Supervision des élèves, enseignants, notes, bulletins, présence et emploi du temps. Pas d’écriture financière.",
                "Contrôler une classe secondaire et vérifier que la trésorerie ne propose aucune écriture.",
            ),
            (
                "Principal secondaire FR\nqa.principal.secondary.fr",
                "Secondaire francophone seulement. Les élèves, classes et enseignants des autres sections doivent rester invisibles.",
                "Comparer ses listes avec celles du principal secondaire global.",
            ),
            (
                "Principal multi-parcours\nqa.principal.multi",
                "Primaire et secondaire uniquement. La maternelle n’est pas incluse, même si le compte dirige plusieurs parcours.",
                "Vérifier les sélecteurs : primaire et secondaire présents, maternelle absente.",
            ),
            (
                "Principal sans affectation\nqa.principal.none",
                "Aucun parcours affecté. Le compte doit être arrêté à l’écran de choix du parcours et ne doit pas entrer dans l’application.",
                "Tester qu’un lien direct vers Élèves ou Finance ne contourne pas le blocage.",
            ),
        ],
    ),
    (
        "Enseignants : primaire, bilingue et emploi du temps",
        [
            (
                "Enseignant primaire FR\nqa.primary.fr",
                "Classe titulaire CE1 A et cohorte bilingue liée. Voit les mêmes 46 élèves dans la projection francophone. Pas d’accès aux autres classes ni à l’inscription.",
                "Tester élèves, appel, notes et emploi du temps de CE1 A.",
            ),
            (
                "Enseignant primaire EN\nqa.primary.en",
                "Projection anglophone Class 3 A de la même cohorte bilingue. Les 46 élèves doivent correspondre au compte primaire FR.",
                "Comparer la liste avec qa.primary.fr et vérifier les noms identiques.",
            ),
            (
                "Enseignant planning FR\nqa.timetable.fr",
                "Accès limité au travail d’emploi du temps qui lui est affecté dans la partie francophone.",
                "Vérifier que seuls ses créneaux et cours autorisés sont exploitables.",
            ),
            (
                "Enseignant planning EN\nqa.timetable.en",
                "Même scénario que le compte précédent, mais dans la partie anglophone.",
                "Comparer les créneaux FR et EN sans voir d’affectation étrangère.",
            ),
        ],
    ),
    (
        "Enseignants : secondaire",
        [
            (
                "Enseignant matière\nqa.sec.subject",
                "6ème A et matière française affectée uniquement. Peut travailler sur sa feuille de notes, ses occurrences de présence et son cahier de textes. Ne génère pas le bulletin de classe.",
                "Essayer de modifier une autre matière et de générer le bulletin : les deux doivent être refusés.",
            ),
            (
                "Autre enseignant matière\nqa.sec.other",
                "6ème A avec ses propres matières : anglais, arts, physique, sciences et SVT selon les affectations de test. Ne doit pas modifier le français.",
                "Comparer les matières visibles avec qa.sec.subject.",
            ),
            (
                "Titulaire secondaire\nqa.sec.titulaire",
                "Supervision complète de la 6ème A et génération du bulletin. Sa propre matière reste modifiable ; les feuilles et occurrences des collègues sont visibles mais en lecture seule.",
                "Ouvrir une matière de collègue, vérifier les champs désactivés, puis tester sa propre matière.",
            ),
        ],
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color: str = GRID, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_run(run, size: float, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_callout(doc: Document, title: str, text: str, fill: str = PALE_GOLD) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, GOLD if fill == PALE_GOLD else TEAL, "8")
    set_cell_margins(cell, 120, 150, 120, 150)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    format_run(paragraph.add_run(f"{title}  "), 10.5, NAVY, True)
    format_run(paragraph.add_run(text), 10.5, INK)


def add_account_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3100, 3900, 2360])
    header = table.rows[0]
    repeat_header(header)
    for cell, label in zip(header.cells, ("Compte", "Ce qui le distingue", "Scénario conseillé")):
        set_cell_shading(cell, NAVY)
        set_cell_borders(cell, WHITE, "4")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        format_run(paragraph.add_run(label), 9.5, WHITE, True)

    for index, (account, difference, scenario) in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        fill = WHITE if index % 2 == 0 else "F5F8FB"
        for cell in row.cells:
            set_cell_shading(cell, fill)
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        role, username = account.split("\n", 1)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        format_run(p.add_run(role), 9.5, NAVY, True)
        p = row.cells[0].add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        format_run(p.add_run(username), 9, TEAL, True)

        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        format_run(p.add_run(difference), 9, INK)

        p = row.cells[2].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        format_run(p.add_run(scenario), 9, INK)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    format_run(header.add_run("BBC SMS  •  COMPTES DE TEST"), 8.5, MUTED, True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    format_run(footer.add_run("BBC Working - environnement de test uniquement"), 8, MUTED)


def add_cover(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(36)
    kicker.paragraph_format.space_after = Pt(8)
    format_run(kicker.add_run("RÉFÉRENCE DE TEST"), 10, GOLD, True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    format_run(title.add_run("Comptes de test BBC SMS"), 27, NAVY, True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    format_run(
        subtitle.add_run("Identifiants, périmètres et scénarios de comparaison"),
        14,
        TEAL,
        False,
    )

    metadata = doc.add_table(rows=3, cols=2)
    set_table_geometry(metadata, [2700, 6660])
    values = (
        ("Environnement", "BBC Working / test - jamais bbc-production"),
        ("Page de connexion", "Ouvrir la page /app/login de l’environnement de test"),
        ("Mot de passe commun", QA_PASSWORD),
    )
    for row, (label, value) in zip(metadata.rows, values):
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_borders(cell)
            set_cell_margins(cell, 100, 140, 100, 140)
        set_cell_shading(row.cells[0], PALE_BLUE)
        set_cell_shading(row.cells[1], WHITE)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        format_run(p.add_run(label), 10, NAVY, True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        format_run(p.add_run(value), 10, INK, label == "Mot de passe commun")

    doc.add_paragraph()
    add_callout(
        doc,
        "IMPORTANT",
        "Ces comptes contiennent des données de démonstration. Ne réutilisez jamais ce mot de passe pour un compte réel et vérifiez toujours que vous êtes sur l’environnement de test avant de vous connecter.",
    )

    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Comment utiliser ce document")
    for label, text in (
        ("Comparer", "connectez-vous avec deux comptes successifs pour observer les différences de menus, de classes et d’actions."),
        ("Tester les limites", "essayez également un lien direct ou une action interdite ; un menu caché ne suffit pas à prouver la sécurité."),
        ("Réinitialiser", "déconnectez-vous complètement entre deux comptes afin d’éviter de conserver le contexte du précédent utilisateur."),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        format_run(paragraph.add_run(f"{label}. "), 10.5, TEAL, True)
        format_run(paragraph.add_run(text), 10.5, INK)

    doc.add_page_break()


def add_final_checklist(doc: Document) -> None:
    doc.add_heading("Contrôles communs à refaire", level=1)
    items = (
        ("Navigation", "le compte ne voit que les modules prévus pour son rôle."),
        ("Liens directs", "une URL copiée depuis un autre rôle ne doit jamais élargir les droits."),
        ("Sélecteurs", "les parcours, classes, matières et élèves hors périmètre restent absents."),
        ("Actions", "les boutons de création, modification, validation et suppression correspondent exactement au rôle."),
        ("Données bilingues", "les enseignants FR et EN voient la même cohorte sous la classe de leur programme."),
        ("Finance", "le comptable peut écrire ; le principal ne dispose que des vues financières autorisées en lecture."),
        ("Déconnexion", "la session précédente est bien supprimée avant de changer d’utilisateur."),
    )
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [900, 8460])
    repeat_header(table.rows[0])
    for cell, label in zip(table.rows[0].cells, ("OK", "Résultat attendu")):
        set_cell_shading(cell, NAVY)
        set_cell_borders(cell, WHITE, "4")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if label == "OK" else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        format_run(p.add_run(label), 9.5, WHITE, True)

    for index, (label, text) in enumerate(items):
        row = table.add_row()
        prevent_row_split(row)
        fill = WHITE if index % 2 == 0 else "F5F8FB"
        for cell in row.cells:
            set_cell_shading(cell, fill)
            set_cell_borders(cell)
            set_cell_margins(cell)
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        format_run(p.add_run("□"), 13, TEAL, True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        format_run(p.add_run(f"{label} : "), 10, NAVY, True)
        format_run(p.add_run(text), 10, INK)

    add_callout(
        doc,
        "EN CAS D’ÉCART",
        "Notez le compte utilisé, le parcours actif, la classe ou matière sélectionnée, l’action tentée et joignez une capture d’écran avant de modifier les permissions.",
        PALE_TEAL,
    )


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    for heading, rows in GROUPS:
        doc.add_heading(heading, level=1)
        add_account_table(doc, rows)

    add_final_checklist(doc)
    doc.core_properties.title = "Comptes de test BBC SMS"
    doc.core_properties.subject = "Identifiants et scénarios de comparaison des rôles"
    doc.core_properties.author = "Bayo Bilingual Complex"
    doc.core_properties.comments = "Référence pour l’environnement BBC Working uniquement."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
