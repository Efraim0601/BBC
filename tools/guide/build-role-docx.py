"""Build the seven bilingual BBC SMS role manuals as polished DOCX files.

The manual content comes from ``role_guides.py``, which is also the source for
the role-aware in-app help.  Keeping one content source prevents the web and
print editions from describing different permissions or procedures.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from role_guides import ROLE_GUIDES


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "output" / "role-guides"
LOGO = ROOT / "frontend" / "public" / "bbc-logo.png"

OUTPUT_NAMES = {
    "administrator": "BBC_SMS_ADMINISTRATOR_GUIDE.docx",
    "principal": "BBC_SMS_PRINCIPAL_GUIDE.docx",
    "prefect": "BBC_SMS_PREFECT_GUIDE.docx",
    "accountant": "BBC_SMS_ACCOUNTANT_GUIDE.docx",
    "primary-teacher": "BBC_SMS_PRIMARY_KINDERGARTEN_TEACHER_GUIDE.docx",
    "secondary-teacher": "BBC_SMS_SECONDARY_TEACHER_GUIDE.docx",
    "parent": "BBC_SMS_PARENT_GUARDIAN_GUIDE.docx",
}

NAVY = "173653"
BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
GOLD = "C89628"
INK = "1A2B3D"
MUTED = "5F7088"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF5D9"
PALE_RED = "FDECEE"
RED = "A82039"
WHITE = "FFFFFF"
LIGHT_BORDER = "C9D4E2"
USABLE_WIDTH_DXA = 9360


def txt(value: dict[str, str], lang: str) -> str:
    return value[lang]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, left: int = 120, right: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, amount in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(amount))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges: dict[str, str | int]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge in edges.items():
        tag = f"w:{edge_name}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge:
                element.set(qn(f"w:{key}"), str(edge[key]))


def set_table_borders(table, color: str = LIGHT_BORDER, size: int = 6) -> None:
    edge = {"val": "single", "sz": size, "space": 0, "color": color}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=edge, bottom=edge, start=edge, end=edge, insideH=edge, insideV=edge)


def configure_fixed_table(table, widths: list[int], *, indent: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid_cols = table._tbl.tblGrid.gridCol_lst
    for index, width in enumerate(widths):
        if index < len(grid_cols):
            grid_cols[index].set(qn("w:w"), str(width))
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def keep_paragraph(paragraph, *, with_next: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext" if with_next else "w:keepLines")
    p_pr.append(keep)


def set_bottom_border(paragraph, color: str, size: int = 14) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_repeatable_numbering(doc: Document) -> None:
    """Install a compact decimal list definition used by the guide preface."""
    numbering = doc.part.numbering_part.element
    abstract_id = max([int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))] or [0]) + 1
    num_id = max([int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))] or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, indent])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    doc._bbc_numbering_id = num_id


def add_numbered_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(doc._bbc_numbering_id))
    num_pr.extend([ilvl, num_id])
    paragraph._p.get_or_add_pPr().append(num_pr)
    paragraph.add_run(text)
    return paragraph


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    number = OxmlElement("w:t")
    number.text = "2"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, number, end])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(10)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.italic = False
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DEEP_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    list_style = styles["List Bullet"]
    list_style.font.name = "Calibri"
    list_style.font.size = Pt(10.5)
    list_style.paragraph_format.left_indent = Inches(0.375)
    list_style.paragraph_format.first_line_indent = Inches(-0.188)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.15


def configure_document(doc: Document, guide: dict) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False

    header_p = section.header.paragraphs[0]
    header_p.text = f"BBC SMS  •  {txt(guide['title'], 'en')}"
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header_p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)

    add_page_number(section.footer.paragraphs[0])

    props = doc.core_properties
    props.title = f"BBC SMS — {txt(guide['title'], 'en')}"
    props.subject = "Bilingual role operating manual"
    props.author = "Bayo Bilingual Complex"
    props.keywords = "BBC SMS, permissions, procedures, role guide"
    props.comments = "Generated from the tested role guide source on 2026-08-28."


def add_cover(doc: Document, guide: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    p.add_run().add_picture(str(LOGO), width=Inches(0.82))

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("BBC SMS  •  MODE D’EMPLOI / OPERATOR GUIDE")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(GOLD)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(txt(guide["title"], "fr"))
    title.add_run("\n")
    english = title.add_run(txt(guide["title"], "en"))
    english.font.size = Pt(22)
    english.font.color.rgb = RGBColor.from_string(DEEP_BLUE)

    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(txt(guide["subtitle"], "fr"))
    sub.add_run("\n")
    sub.add_run(txt(guide["subtitle"], "en"))

    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider.paragraph_format.left_indent = Inches(1.45)
    divider.paragraph_format.right_indent = Inches(1.45)
    set_bottom_border(divider, GOLD, 18)

    scope = doc.add_table(rows=2, cols=1)
    configure_fixed_table(scope, [USABLE_WIDTH_DXA])
    set_table_borders(scope, color=NAVY, size=8)
    set_cell_shading(scope.cell(0, 0), NAVY)
    head = scope.cell(0, 0).paragraphs[0]
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = head.add_run("PÉRIMÈTRE / SCOPE")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)
    body = scope.cell(1, 0).paragraphs[0]
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body.add_run(txt(guide["scope"], "fr") + "\n").bold = True
    e = body.add_run(txt(guide["scope"], "en"))
    e.italic = True
    e.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(18)
    meta.add_run("Version vérifiée / Verified edition\n").bold = True
    meta.add_run("28 août 2026 / 28 August 2026\n")
    meta.add_run("Application locale testée sur ports 8130 / 8131")
    for run in meta.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_page_break()


def add_bilingual_cell(cell, fr: str, en: str, *, label: str | None = None) -> None:
    paragraph = cell.paragraphs[0]
    if label:
        run = paragraph.add_run(label + "\n")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    fr_run = paragraph.add_run(fr + "\n")
    fr_run.bold = True
    en_run = paragraph.add_run(en)
    en_run.italic = True
    en_run.font.color.rgb = RGBColor.from_string(MUTED)


def add_section_heading(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph(title, style="Heading 1")
    set_bottom_border(p, GOLD, 10)
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.paragraph_format.space_after = Pt(10)
        for run in sub.runs:
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(MUTED)


def add_overview(doc: Document, guide: dict) -> None:
    add_section_heading(doc, "Le rôle en bref / Role at a glance")

    overview = doc.add_table(rows=2, cols=1)
    configure_fixed_table(overview, [USABLE_WIDTH_DXA])
    set_table_borders(overview)
    set_cell_shading(overview.cell(0, 0), PALE_BLUE)
    add_bilingual_cell(
        overview.cell(0, 0),
        txt(guide["summary"], "fr"),
        txt(guide["summary"], "en"),
        label="MISSION / PURPOSE",
    )
    add_bilingual_cell(
        overview.cell(1, 0),
        txt(guide["scope"], "fr"),
        txt(guide["scope"], "en"),
        label="PÉRIMÈTRE / SCOPE",
    )

    doc.add_heading("Comment utiliser ce guide / How to use this guide", level=2)
    instructions = [
        "Vérifiez le rôle et le parcours affichés dans l’en-tête. / Check the role and parcours shown in the header.",
        "Suivez la procédure qui correspond à l’action à réaliser. / Follow the workflow for the task you need to perform.",
        "Respectez les limites, même si un bouton inattendu apparaît. / Respect the boundaries even if an unexpected button appears.",
        "Terminez par les contrôles rapides de la dernière section. / Finish with the quick checks in the final section.",
    ]
    for instruction in instructions:
        add_numbered_paragraph(doc, instruction)


def add_permissions(doc: Document, guide: dict) -> None:
    add_section_heading(
        doc,
        "Ce que ce rôle peut faire / What this role can do",
        "Les autorisations restent limitées au périmètre indiqué ci-dessus. / Permissions remain limited to the scope shown above.",
    )
    table = doc.add_table(rows=1, cols=2)
    configure_fixed_table(table, [4680, 4680])
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for index, label in enumerate(("FRANÇAIS", "ENGLISH")):
        set_cell_shading(table.cell(0, index), NAVY)
        p = table.cell(0, index).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    for name, description in guide["permissions"]:
        row = table.add_row()
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell)
        add_bilingual_permission(row.cells[0], txt(name, "fr"), txt(description, "fr"), BLUE)
        add_bilingual_permission(row.cells[1], txt(name, "en"), txt(description, "en"), DEEP_BLUE)


def add_bilingual_permission(cell, title: str, body: str, color: str) -> None:
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(body)


def add_workflows(doc: Document, guide: dict) -> None:
    add_section_heading(
        doc,
        "Procédures pas à pas / Step-by-step workflows",
        "Les noms de pages et les routes correspondent à la version locale vérifiée. / Page names and routes match the verified local build.",
    )
    for workflow_number, workflow in enumerate(guide["workflows"], start=1):
        heading = doc.add_paragraph(style="Heading 2")
        if workflow.get("page_break_before"):
            heading.paragraph_format.page_break_before = True
        heading.add_run(f"{workflow_number}. {txt(workflow['title'], 'fr')}")
        english = heading.add_run(f" / {txt(workflow['title'], 'en')}")
        english.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
        keep_paragraph(heading, with_next=True)

        route = doc.add_paragraph()
        route.paragraph_format.space_after = Pt(5)
        route.paragraph_format.keep_with_next = True
        r = route.add_run("PAGE / ROUTE  ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(GOLD)
        code = route.add_run(workflow.get("route") or "—")
        code.font.name = "Consolas"
        code.font.size = Pt(9)
        code.font.color.rgb = RGBColor.from_string(NAVY)

        steps = workflow["steps"]
        table = doc.add_table(rows=1, cols=3)
        configure_fixed_table(table, [720, 4320, 4320])
        set_table_borders(table)
        repeat_table_header(table.rows[0])
        for index, label in enumerate(("#", "FRANÇAIS", "ENGLISH")):
            set_cell_shading(table.cell(0, index), PALE_BLUE)
            p = table.cell(0, index).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run(label)
            rr.bold = True
            rr.font.color.rgb = RGBColor.from_string(NAVY)
        for step_number, step in enumerate(steps, start=1):
            row = table.add_row()
            prevent_row_split(row)
            for cell in row.cells:
                set_cell_margins(cell)
            num_p = row.cells[0].paragraphs[0]
            num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            num = num_p.add_run(str(step_number))
            num.bold = True
            num.font.color.rgb = RGBColor.from_string(GOLD)
            row.cells[1].paragraphs[0].add_run(txt(step, "fr"))
            en = row.cells[2].paragraphs[0].add_run(txt(step, "en"))
            en.font.color.rgb = RGBColor.from_string(MUTED)

        if workflow.get("note"):
            note = doc.add_table(rows=1, cols=1)
            configure_fixed_table(note, [USABLE_WIDTH_DXA])
            set_cell_shading(note.cell(0, 0), PALE_GOLD)
            set_cell_border(
                note.cell(0, 0),
                start={"val": "single", "sz": 28, "space": 0, "color": GOLD},
                top={"val": "nil"},
                bottom={"val": "nil"},
                end={"val": "nil"},
            )
            add_bilingual_cell(
                note.cell(0, 0),
                txt(workflow["note"], "fr"),
                txt(workflow["note"], "en"),
                label="IMPORTANT",
            )


def add_parallel_list(doc: Document, heading: str, items: Iterable[dict[str, str]], *, warning: bool = False) -> None:
    doc.add_heading(heading, level=2)
    table = doc.add_table(rows=1, cols=3)
    configure_fixed_table(table, [720, 4320, 4320])
    set_table_borders(table, color=RED if warning else LIGHT_BORDER)
    repeat_table_header(table.rows[0])
    fill = PALE_RED if warning else PALE_BLUE
    for index, label in enumerate(("", "FRANÇAIS", "ENGLISH")):
        set_cell_shading(table.cell(0, index), fill)
        p = table.cell(0, index).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(RED if warning else NAVY)
    for index, item in enumerate(items, start=1):
        row = table.add_row()
        prevent_row_split(row)
        marker = row.cells[0].paragraphs[0]
        marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = marker.add_run("!" if warning else "✓")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(RED if warning else GOLD)
        row.cells[1].paragraphs[0].add_run(txt(item, "fr"))
        en = row.cells[2].paragraphs[0].add_run(txt(item, "en"))
        en.font.color.rgb = RGBColor.from_string(MUTED)


def add_boundaries_and_checks(doc: Document, guide: dict) -> None:
    add_section_heading(doc, "Limites et contrôles / Boundaries and checks")
    add_parallel_list(doc, "Ce que ce rôle ne doit pas faire / What this role must not do", guide["boundaries"])
    add_parallel_list(doc, "Contrôle rapide avant de terminer / Quick verification before finishing", guide["verification"])

    if guide["known_gaps"]:
        add_parallel_list(
            doc,
            "Écarts confirmés dans la version testée / Confirmed gaps in the tested build",
            guide["known_gaps"],
            warning=True,
        )
        warning = doc.add_paragraph()
        warning.paragraph_format.space_before = Pt(8)
        run = warning.add_run(
            "Ces écarts ne constituent pas des autorisations. Respectez le modèle décrit dans ce guide et signalez le comportement à l’administrateur.\n"
        )
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(RED)
        en = warning.add_run(
            "These gaps are not permissions. Follow the operating model in this guide and report the behavior to the administrator."
        )
        en.italic = True
        en.font.color.rgb = RGBColor.from_string(RED)
        keep_paragraph(warning)

    doc.add_heading("Aide / Help", level=2)
    help_p = doc.add_paragraph()
    help_p.add_run("Dans l’application / In the application: ").bold = True
    help_p.add_run("cliquez sur Help dans l’en-tête; le rôle connecté ouvre automatiquement ce guide. ")
    help_p.add_run("Click Help in the header; the signed-in role automatically opens this guide.")
    help_p.add_run("\n")
    help_p.add_run("Avant de demander une modification d’accès / Before requesting an access change: ").bold = True
    help_p.add_run("notez le rôle, le parcours, la classe, la page, l’action et l’heure du problème. / Record the role, parcours, class, page, action, and time of the issue.")

def build_one(guide: dict) -> Path:
    doc = Document()
    configure_styles(doc)
    configure_document(doc, guide)
    set_repeatable_numbering(doc)
    add_cover(doc, guide)
    add_overview(doc, guide)
    add_permissions(doc, guide)
    add_workflows(doc, guide)
    add_boundaries_and_checks(doc, guide)

    output = OUT / OUTPUT_NAMES[guide["slug"]]
    doc.save(output)
    return output


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    outputs = [build_one(guide) for guide in ROLE_GUIDES]
    for output in outputs:
        print(f"Built {output.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
