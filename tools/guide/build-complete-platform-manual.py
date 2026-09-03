from __future__ import annotations

import json
import os
from dataclasses import dataclass
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SCREENSHOTS = ROOT / "output" / "platform-user-manual" / "screenshots"
MANIFEST = SCREENSHOTS / "manifest.json"
OUT = ROOT / "output" / "platform-user-manual" / "GUIDE_COMPLET_BBC_SMS_FR_2026-08-31.docx"

NAVY = "163653"
TEAL = "0F8278"
GOLD = "D8A928"
INK = "172B42"
MUTED = "5B6B80"
PALE = "EEF4F8"
PALE_TEAL = "E8F7F4"
WHITE = "FFFFFF"
RED = "A63D40"


@dataclass(frozen=True)
class Task:
    title: str
    goal: str
    steps: tuple[str, ...]
    images: tuple[str, ...]
    tip: str = ""


CHAPTERS: list[tuple[str, str, list[Task]]] = [
    (
        "1. Démarrer et se repérer",
        "Connexion, choix du contexte de travail et navigation générale.",
        [
            Task(
                "Se connecter",
                "Ouvrir une session BBC SMS et atteindre l’espace de travail.",
                (
                    "Ouvrez BBC SMS dans votre navigateur.",
                    "Saisissez votre identifiant et votre mot de passe.",
                    "Cliquez sur « Se connecter ».",
                    "Si un choix de parcours apparaît, sélectionnez le niveau et la section sur lesquels vous devez travailler.",
                ),
                ("000-login.png",),
                "Ne partagez jamais votre mot de passe. Utilisez le bouton de déconnexion avant de quitter un poste partagé.",
            ),
            Task(
                "Ouvrir un module",
                "Trouver rapidement la fonction dont vous avez besoin.",
                (
                    "Cliquez sur « Applications » dans le menu de gauche.",
                    "Repérez la carte du module souhaité : Élèves, Académique, Présence, Finance, etc.",
                    "Cliquez sur la carte. Le même module reste ensuite disponible dans le menu de gauche.",
                ),
                ("001-applications.png",),
                "Les cartes visibles dépendent de votre travail. L’absence d’une carte signifie que la fonction n’est pas proposée dans votre espace.",
            ),
            Task(
                "Lire le tableau de bord",
                "Consulter les principaux indicateurs de l’établissement.",
                (
                    "Cliquez sur « Tableau de bord ».",
                    "Lisez les indicateurs de synthèse en haut de page.",
                    "Utilisez les graphiques et listes d’alertes pour repérer les éléments à traiter.",
                    "Actualisez la page après une opération importante si un chiffre doit être recalculé.",
                ),
                ("002-dashboard.png",),
            ),
        ],
    ),
    (
        "2. Élèves et familles",
        "Rechercher, inscrire, importer et consulter le dossier d’un élève.",
        [
            Task(
                "Rechercher, filtrer et exporter la liste des élèves",
                "Afficher une classe précise, retrouver un élève et produire une liste Excel ou PDF.",
                (
                    "Cliquez sur « Élèves ».",
                    "Choisissez une classe dans le filtre de classe, ou utilisez la recherche par nom ou matricule.",
                    "Vérifiez le nombre de résultats et ouvrez une fiche en cliquant sur la ligne de l’élève.",
                    "Pour obtenir une liste, ouvrez « Exporter » puis choisissez Excel ou PDF.",
                ),
                ("010-students-list.png", "011-student-detail.png"),
                "Les filtres appliqués à l’écran déterminent la liste exportée.",
            ),
            Task(
                "Inscrire un élève",
                "Créer un dossier élève, choisir son parcours et enregistrer sa famille même sans adresse e-mail.",
                (
                    "Dans « Élèves », cliquez sur « Nouvel élève ».",
                    "Renseignez le nom. Le prénom est facultatif.",
                    "Saisissez la date de naissance au format JJ/MM/AAAA ; les barres obliques s’ajoutent automatiquement sur mobile.",
                    "Cliquez sur « Continuer », puis choisissez le niveau, la section et la classe ou cohorte proposée.",
                    "Ajoutez les responsables légaux disponibles. Leur e-mail est facultatif et pourra être ajouté plus tard.",
                    "Relisez le résumé et confirmez l’inscription.",
                ),
                ("012-student-create-identity.png", "013-student-create-schooling.png"),
                "Pour une classe bilingue liée, l’élève appartient à une seule cohorte mais apparaît dans les deux classes programme.",
            ),
            Task(
                "Importer plusieurs élèves et familles",
                "Préparer et contrôler un import de plusieurs dossiers en une seule opération.",
                (
                    "Dans « Élèves », ouvrez l’outil d’import des élèves et familles.",
                    "Téléchargez le modèle proposé et complétez une ligne par élève.",
                    "Conservez les colonnes du modèle. Les e-mails parentaux peuvent rester vides.",
                    "Déposez le fichier, lancez la prévisualisation et corrigez les lignes signalées.",
                    "Validez uniquement lorsque le résumé correspond au nombre attendu d’élèves et de responsables.",
                ),
                ("014-student-family-import.png",),
                "Commencez par un petit fichier de deux ou trois élèves lorsque vous utilisez un nouveau modèle.",
            ),
        ],
    ),
    (
        "3. Parcours scolaire, promotion, santé et documents",
        "Suivre la vie scolaire d’un élève au-delà de sa fiche d’identité.",
        [
            Task(
                "Consulter le parcours scolaire",
                "Voir l’historique des inscriptions, résultats publiés et décisions de passage.",
                (
                    "Cliquez sur « Parcours ».",
                    "Choisissez la classe puis l’élève.",
                    "Parcourez la chronologie des sessions, classes et décisions officielles.",
                    "Utilisez cet écran pour vérifier une incohérence avant toute promotion ou réinscription.",
                ),
                ("020-journey.png",),
            ),
            Task(
                "Préparer et suivre un passage de classe",
                "Définir les décisions de fin d’année et préparer la session suivante.",
                (
                    "Cliquez sur « Passage de classe ».",
                    "Choisissez la session et la classe source.",
                    "Contrôlez les règles, les résultats disponibles et la destination proposée.",
                    "Enregistrez les décisions en brouillon, relisez-les, puis utilisez la clôture uniquement lorsque tout est validé.",
                    "Ouvrez le registre de promotion pour vérifier l’historique et la traçabilité des opérations.",
                ),
                ("021-promotion.png", "023-promotion-register.png"),
                "La clôture annuelle est une étape importante : vérifiez les effectifs et les destinations avant de confirmer.",
            ),
            Task(
                "Choisir manuellement le parcours de la prochaine session",
                "Orienter un élève vers le parcours francophone ou anglophone lorsque les cohortes se séparent.",
                (
                    "Ouvrez « Orientation » depuis le module de passage de classe.",
                    "Sélectionnez la session source, la classe ou cohorte et l’élève.",
                    "Choisissez manuellement le parcours de la session suivante.",
                    "Enregistrez et vérifiez la décision dans le registre de promotion.",
                ),
                ("022-pathways.png",),
                "Le choix est volontairement manuel d’une session à l’autre ; le système ne décide pas à la place de l’école.",
            ),
            Task(
                "Tenir le dossier de santé",
                "Consulter les informations médicales utiles et enregistrer un passage à l’infirmerie.",
                (
                    "Cliquez sur « Santé ».",
                    "Choisissez la classe puis l’élève.",
                    "Consultez les allergies, informations importantes et visites existantes.",
                    "Pour une nouvelle visite, renseignez la date, le motif, les observations et la suite donnée, puis enregistrez.",
                ),
                ("024-health.png",),
                "Ne consignez que les informations nécessaires au suivi scolaire et protégez la confidentialité de l’élève.",
            ),
            Task(
                "Gérer les documents d’un élève",
                "Déposer, retrouver, consulter et télécharger les pièces du dossier élève.",
                (
                    "Cliquez sur « Documents ».",
                    "Choisissez la classe puis l’élève.",
                    "Sélectionnez la catégorie du document et ajoutez un libellé clair.",
                    "Déposez le fichier, enregistrez, puis vérifiez qu’il apparaît dans la liste.",
                    "Cliquez sur un document pour le consulter ou utilisez le téléchargement.",
                ),
                ("025-student-documents.png",),
            ),
        ],
    ),
    (
        "4. Personnel",
        "Créer et suivre les employés, leurs documents, candidatures, congés et rémunérations.",
        [
            Task(
                "Rechercher et ouvrir une fiche employé",
                "Utiliser l’annuaire et consulter les informations professionnelles d’un employé.",
                (
                    "Cliquez sur « Personnel ».",
                    "Utilisez la recherche et les filtres de rôle, département ou statut.",
                    "Cliquez sur la ligne d’un employé pour ouvrir sa fiche dédiée.",
                    "Cliquez sur « Modifier » pour corriger les coordonnées, l’affectation, la rémunération ou les pièces jointes.",
                ),
                ("030-staff-list.png", "031-staff-detail.png", "032-staff-edit.png"),
            ),
            Task(
                "Consulter les documents du personnel",
                "Voir un CV, diplôme, pièce d’identité, certificat, contrat ou autre document sans quitter la fiche.",
                (
                    "Ouvrez la fiche de l’employé.",
                    "Dans « Documents du personnel », cliquez sur le nom du document.",
                    "Consultez l’image ou le PDF directement dans la page.",
                    "Utilisez « Télécharger » si vous devez conserver une copie locale.",
                ),
                ("031b-staff-document-preview.png",),
            ),
            Task(
                "Créer un employé et joindre plusieurs catégories de documents",
                "Créer une fiche complète avec plusieurs pièces différentes lors de la même saisie.",
                (
                    "Dans « Personnel », cliquez sur « Nouvel employé ».",
                    "Renseignez l’identité, le type de contrat, les coordonnées, le rôle professionnel et la rémunération.",
                    "Choisissez les classes ou niveaux utiles à l’affectation.",
                    "Dans la zone des documents, choisissez une catégorie, ajoutez le fichier, puis cliquez sur « Ajouter une catégorie » pour chaque autre type de pièce.",
                    "Enregistrez et ouvrez la nouvelle fiche pour vérifier toutes les pièces.",
                ),
                ("033-staff-create.png", "034-staff-document-categories.png"),
            ),
            Task(
                "Traiter les candidatures",
                "Publier le formulaire de candidature, examiner une demande et convertir la personne retenue en employé.",
                (
                    "Dans « Personnel », ouvrez l’onglet « Candidatures ».",
                    "Activez le portail public lorsque l’école souhaite recevoir des candidatures et copiez le lien proposé.",
                    "Le candidat ouvre le lien, complète le formulaire et envoie sa demande.",
                    "Dans la liste interne, ouvrez la demande, puis choisissez Accepter, Refuser ou Finaliser.",
                    "Lors de la finalisation, complétez le contrat, le salaire et l’affectation avant de créer la fiche employé.",
                ),
                ("035-staff-applications.png", "039-public-staff-application.png"),
            ),
            Task(
                "Gérer les départements",
                "Créer l’organisation interne et désigner les responsables de département.",
                (
                    "Dans « Personnel », ouvrez « Départements ».",
                    "Cliquez sur le bouton de création, saisissez le nom et choisissez éventuellement un responsable.",
                    "Enregistrez, puis modifiez ou retirez un département depuis sa ligne.",
                ),
                ("036-staff-departments.png",),
            ),
            Task(
                "Gérer les congés et la masse salariale",
                "Enregistrer une demande de congé et contrôler la synthèse des rémunérations.",
                (
                    "Ouvrez l’onglet « Congés », puis créez une demande avec l’employé, le type, les dates et le motif.",
                    "Examinez la demande et enregistrez la décision.",
                    "Ouvrez ensuite « Masse salariale » pour consulter les montants mensuels et annuels consolidés.",
                    "Corrigez une rémunération depuis la fiche employé si la synthèse est incorrecte.",
                ),
                ("037-staff-leave.png", "038-staff-salary.png"),
            ),
        ],
    ),
    (
        "5. Académique",
        "Saisir les notes, préparer les conseils, contrôler les résultats et produire les bulletins.",
        [
            Task(
                "Consulter ou préparer un bulletin",
                "Choisir une classe, une période et un élève pour afficher le bulletin correspondant.",
                (
                    "Cliquez sur « Académique » puis sur l’onglet « Bulletin ».",
                    "Choisissez la classe et le jalon académique.",
                    "Sélectionnez un élève dans la liste.",
                    "Contrôlez les matières, moyennes, rang, assiduité et observations avant toute génération officielle.",
                ),
                ("040-academic-report-card-roster.png",),
            ),
            Task(
                "Saisir les notes d’une classe",
                "Remplir une feuille de notes complète et la transmettre pour validation.",
                (
                    "Dans « Académique », cliquez sur « Saisie des notes ».",
                    "Choisissez la classe, la période et la matière.",
                    "Vérifiez que le nom de la matière et le nombre d’élèves sont corrects.",
                    "Saisissez une note pour chaque élève, ou choisissez le statut prévu lorsqu’une note ne s’applique pas.",
                    "Ajoutez les appréciations utiles, enregistrez le brouillon, puis envoyez la feuille lorsque la saisie est complète.",
                ),
                ("041-academic-grade-entry.png",),
                "La capture montre volontairement les lignes d’élèves, pas seulement les filtres du haut de page.",
            ),
            Task(
                "Renseigner l’assiduité et le conseil de classe",
                "Calculer automatiquement les absences sur une période et compléter les décisions du conseil.",
                (
                    "Ouvrez « Assiduité & conseil ».",
                    "Choisissez la classe et la séquence.",
                    "Indiquez l’intervalle de dates ou utilisez toute la séquence, puis cliquez sur « Appliquer ».",
                    "Les absences non justifiées sont calculées automatiquement ; ouvrez un élève avec le bouton + pour compléter les informations nécessaires.",
                    "Enregistrez les appréciations et décisions du conseil uniquement après contrôle des données de présence.",
                ),
                ("042-academic-council.png",),
            ),
            Task(
                "Contrôler la classe et le procès-verbal",
                "Comparer les résultats de toute la classe avant validation et publication.",
                (
                    "Ouvrez « Vue de classe » pour vérifier les moyennes, rangs, données manquantes et anomalies.",
                    "Corrigez les feuilles concernées avant de continuer.",
                    "Ouvrez ensuite « Procès-verbal » et cliquez sur « Charger ».",
                    "Contrôlez le tableau consolidé et utilisez l’export ou l’impression lorsque le résultat est complet.",
                ),
                ("043-academic-class-overview.png", "044-academic-master-sheet.png"),
            ),
            Task(
                "Générer plusieurs bulletins",
                "Lancer et suivre une génération de bulletins pour toute une classe.",
                (
                    "Ouvrez « Génération en lot ».",
                    "Choisissez la classe, la période et le programme de bulletin.",
                    "Lancez d’abord le contrôle de préparation et corrigez les blocages signalés.",
                    "Démarrez la génération, suivez le nombre de réussites et d’échecs, puis téléchargez les documents terminés.",
                ),
                ("045-academic-batch.png",),
            ),
        ],
    ),
    (
        "6. Présence",
        "Faire l’appel, analyser l’assiduité et rapprocher les pointages.",
        [
            Task(
                "Faire l’appel quotidien ou par période",
                "Marquer rapidement les présents, absents, retards et excusés.",
                (
                    "Cliquez sur « Présence » puis « Liste d’appel ».",
                    "Choisissez la date et la classe. Au secondaire, choisissez aussi la période publiée.",
                    "Utilisez « Tous présents » puis corrigez uniquement les exceptions, ou marquez les statuts élève par élève.",
                    "Le motif d’absence est facultatif. Ajoutez-le uniquement lorsque l’information est connue.",
                    "Enregistrez, relisez les totaux, puis finalisez pour verrouiller l’appel.",
                ),
                ("050-attendance-daily.png",),
            ),
            Task(
                "Analyser la présence",
                "Repérer les absences répétées et les tendances par période.",
                (
                    "Dans « Présence », ouvrez « Analyses ».",
                    "Choisissez la période, la classe et les filtres souhaités.",
                    "Examinez les indicateurs, puis ouvrez les élèves nécessitant un suivi.",
                    "Exportez les résultats si une réunion ou un contrôle l’exige.",
                ),
                ("051-attendance-analytics.png",),
            ),
            Task(
                "Rapprocher les lecteurs de présence",
                "Contrôler les lecteurs biométriques et rapprocher leurs pointages avec les dossiers.",
                (
                    "Ouvrez « Lecteurs & rapprochement ».",
                    "Vérifiez l’état et la dernière communication de chaque lecteur.",
                    "Consultez les pointages non rapprochés et associez-les à la bonne personne.",
                    "Relancez le rapprochement après correction et contrôlez les exceptions restantes.",
                ),
                ("052-attendance-devices.png",),
            ),
        ],
    ),
    (
        "7. Discipline et cahier de textes",
        "Consigner les incidents et assurer le suivi pédagogique quotidien.",
        [
            Task(
                "Consulter et enregistrer un incident",
                "Créer une trace structurée d’un incident et suivre les actions associées.",
                (
                    "Cliquez sur « Discipline ».",
                    "Utilisez les filtres pour retrouver un incident existant.",
                    "Cliquez sur « Nouvel incident ».",
                    "Choisissez l’élève, la date, la catégorie, la gravité et décrivez les faits avec précision.",
                    "Ajoutez la mesure prise et les notifications nécessaires, puis enregistrez.",
                ),
                ("060-discipline-list.png", "061-discipline-create.png"),
            ),
            Task(
                "Tenir le cahier de textes",
                "Enregistrer ce qui a été fait en classe et les devoirs à venir.",
                (
                    "Cliquez sur « Cahier de textes ».",
                    "Choisissez la classe et la matière lorsqu’elle est demandée.",
                    "Consultez les entrées précédentes pour éviter les doublons.",
                    "Cliquez sur « Nouvelle entrée », indiquez la date, le contenu du cours, les devoirs et l’échéance, puis enregistrez.",
                ),
                ("062-coursebook.png", "063-coursebook-create.png"),
            ),
        ],
    ),
    (
        "8. Emploi du temps",
        "Construire, publier et consulter les plannings de classes, enseignants et salles.",
        [
            Task(
                "Construire le planning d’une classe",
                "Configurer une classe et remplir sa grille hebdomadaire.",
                (
                    "Cliquez sur « Emploi du temps » puis « Planning des classes ».",
                    "Choisissez la classe et vérifiez le modèle appliqué.",
                    "Cliquez sur une case vide pour ajouter la matière, l’enseignant et la salle.",
                    "Corrigez les conflits signalés, relisez toute la semaine, puis publiez la classe.",
                    "Utilisez « Rouvrir cette classe » uniquement lorsqu’une modification du planning publié est réellement nécessaire.",
                ),
                ("064-timetable-class.png",),
            ),
            Task(
                "Utiliser la vue maître",
                "Contrôler l’ensemble des classes et détecter les conflits globaux.",
                (
                    "Ouvrez l’onglet « Vue maître ».",
                    "Choisissez le jour et les filtres utiles.",
                    "Repérez les chevauchements d’enseignant ou de salle.",
                    "Revenez au planning de la classe concernée pour corriger, puis rechargez la vue maître.",
                ),
                ("065-timetable-master.png",),
            ),
            Task(
                "Consulter le planning d’un enseignant ou d’une salle",
                "Vérifier les disponibilités et éviter une double occupation.",
                (
                    "Ouvrez « Planning des enseignants » et choisissez l’enseignant.",
                    "Contrôlez sa semaine, les classes et les matières associées.",
                    "Ouvrez ensuite « Salles » et choisissez une salle pour vérifier son occupation.",
                    "Corrigez les conflits depuis le planning de classe.",
                ),
                ("066-timetable-teachers.png", "067-timetable-rooms.png"),
            ),
            Task(
                "Enregistrer un remplacement",
                "Remplacer temporairement un enseignant sur une occurrence publiée.",
                (
                    "Ouvrez « Remplacements ».",
                    "Choisissez la date et l’occurrence concernée.",
                    "Sélectionnez l’enseignant remplaçant et indiquez le motif.",
                    "Enregistrez puis vérifiez le planning de l’enseignant et la liste d’appel concernée.",
                ),
                ("068-timetable-substitutions.png",),
            ),
            Task(
                "Configurer les règles et périodes horaires",
                "Définir les disponibilités des enseignants et les créneaux utilisables.",
                (
                    "Ouvrez « Règles enseignants » pour renseigner qualifications, disponibilités et charge maximale.",
                    "Enregistrez les contraintes avant de construire le planning.",
                    "Ouvrez « Périodes horaires » pour définir les heures de début et de fin, les pauses et l’ordre des créneaux.",
                    "Contrôlez qu’aucune période ne se chevauche avant de publier les emplois du temps.",
                ),
                ("069-timetable-rules.png", "070-timetable-periods.png"),
            ),
        ],
    ),
    (
        "9. Communication et ressources",
        "Partager les événements, correspondances, documents et listes de classe.",
        [
            Task(
                "Créer un événement",
                "Publier une annonce ou un événement dans le calendrier scolaire.",
                (
                    "Cliquez sur « Événements ».",
                    "Parcourez le calendrier pour vérifier qu’un événement similaire n’existe pas déjà.",
                    "Cliquez sur « Nouvel événement ».",
                    "Renseignez le titre, les dates, le public, le lieu et les détails, puis publiez.",
                ),
                ("080-events.png", "081-events-create.png"),
            ),
            Task(
                "Échanger avec une famille",
                "Lire le fil d’un élève et envoyer une nouvelle note.",
                (
                    "Cliquez sur « Correspondance ».",
                    "Choisissez la classe puis l’élève.",
                    "Lisez le fil existant et les accusés de réception.",
                    "Cliquez sur « Nouvelle note », rédigez un objet et un message clairs, choisissez si un accusé est requis, puis envoyez.",
                ),
                ("082-correspondence.png", "083-correspondence-create.png"),
            ),
            Task(
                "Partager des ressources",
                "Mettre des documents à disposition des utilisateurs concernés.",
                (
                    "Cliquez sur « Ressources ».",
                    "Choisissez la catégorie et le public concerné.",
                    "Ajoutez un titre explicite, une description et le fichier.",
                    "Enregistrez puis ouvrez la ressource pour vérifier le téléchargement.",
                ),
                ("084-library.png",),
            ),
            Task(
                "Gérer les fournitures et manuels",
                "Préparer les listes demandées à une classe et les rendre visibles aux familles.",
                (
                    "Cliquez sur « Fournitures & livres ».",
                    "Choisissez la classe.",
                    "Dans « Fournitures », ajoutez chaque article, quantité et indication facultative.",
                    "Ouvrez « Manuels scolaires » et renseignez les titres, auteurs, éditeurs et informations utiles.",
                    "Enregistrez et vérifiez le total ainsi que l’ordre d’affichage.",
                ),
                ("085-classkit-supplies.png", "086-classkit-books.png"),
            ),
        ],
    ),
    (
        "10. Configuration fonctionnelle",
        "Préparer les données académiques et générales nécessaires au fonctionnement quotidien.",
        [
            Task(
                "Configurer les sections et classes",
                "Créer la structure des niveaux et des classes de l’établissement.",
                (
                    "Cliquez sur « Paramètres » puis « Configuration académique ».",
                    "Ouvrez « Sections » pour contrôler les parcours francophone et anglophone.",
                    "Ouvrez « Classes », puis ajoutez ou modifiez le nom, le niveau, la section et la capacité.",
                    "Vérifiez les classes avant d’inscrire des élèves ou de construire les emplois du temps.",
                ),
                ("090-settings-sections.png", "091-settings-classes.png"),
            ),
            Task(
                "Configurer les matières par classe",
                "Créer le catalogue de matières, les coefficients et les affectations pédagogiques.",
                (
                    "Ouvrez « Matières » et créez chaque matière avec un code unique et un libellé clair.",
                    "Ouvrez « Matières par classe ».",
                    "Choisissez la session et la classe, ajoutez les matières, les coefficients et les enseignants responsables.",
                    "Enregistrez puis contrôlez la feuille de notes et l’emploi du temps de la classe.",
                ),
                ("092-settings-subjects.png", "093-settings-curriculum.png"),
            ),
            Task(
                "Associer des classes bilingues",
                "Faire apparaître la même cohorte d’élèves dans deux classes programme liées.",
                (
                    "Ouvrez « Associer classes bilingues ».",
                    "Choisissez le niveau maternelle ou primaire concerné.",
                    "Sélectionnez la classe francophone et la classe anglophone correspondantes, par exemple SIL A et Class 1 A.",
                    "Enregistrez le groupe et contrôlez son effectif.",
                    "Vérifiez ensuite la liste des élèves, l’appel, les notes et les deux bulletins programme.",
                ),
                ("094-settings-bilingual.png",),
                "Les enseignants des deux programmes restent différents même si les élèves sont les mêmes.",
            ),
            Task(
                "Configurer les évaluations et coefficients",
                "Définir les colonnes de notes utilisées dans chaque séquence.",
                (
                    "Ouvrez « Évaluations ».",
                    "Choisissez la période et le périmètre demandé.",
                    "Ajoutez les évaluations dans l’ordre souhaité, avec leur barème et coefficient.",
                    "Enregistrez puis vérifiez la feuille de saisie des notes avant le début des évaluations.",
                ),
                ("096-settings-assessments.png",),
            ),
            Task(
                "Configurer les modèles et l’identité visuelle",
                "Préparer l’apparence des bulletins, reçus et autres documents officiels.",
                (
                    "Ouvrez « Modèles / marque ».",
                    "Ajoutez ou remplacez le logo et les informations d’en-tête.",
                    "Choisissez le type de document et modifiez le modèle correspondant.",
                    "Générez un aperçu et relisez toutes les informations avant d’activer la nouvelle version.",
                ),
                ("097-settings-branding.png",),
            ),
            Task(
                "Configurer les années, trimestres et séquences",
                "Créer le calendrier académique utilisé par les notes, bulletins, présences et promotions.",
                (
                    "Ouvrez « Années & périodes ».",
                    "Créez ou sélectionnez l’année scolaire et définissez ses dates.",
                    "Ajoutez les trimestres et séquences dans l’ordre, sans chevauchement.",
                    "Définissez les fenêtres de saisie et de validation, puis contrôlez la checklist de préparation.",
                ),
                ("098-settings-sessions.png",),
            ),
            Task(
                "Configurer les informations générales, le calendrier et la discipline",
                "Renseigner les données de l’école et les catalogues utilisés au quotidien.",
                (
                    "Ouvrez « Général » et complétez le nom, l’adresse et les coordonnées officielles.",
                    "Ouvrez « Calendrier » et ajoutez jours fériés, vacances et journées spéciales.",
                    "Ouvrez « Discipline » et configurez les catégories, niveaux de gravité et mesures disponibles.",
                    "Enregistrez chaque écran puis vérifiez les modules qui utilisent ces données.",
                ),
                ("099-settings-general.png", "100-settings-calendar.png", "101-settings-discipline.png"),
            ),
        ],
    ),
    (
        "11. Pilotage et rapports",
        "Repérer les situations à traiter et produire des synthèses.",
        [
            Task(
                "Traiter les alertes",
                "Identifier les risques scolaires, administratifs ou financiers nécessitant une action.",
                (
                    "Cliquez sur « Alertes ».",
                    "Filtrez par type, priorité ou statut.",
                    "Ouvrez une alerte, consultez les éléments qui l’ont déclenchée et attribuez l’action à mener.",
                    "Mettez à jour le statut lorsque le suivi est terminé.",
                ),
                ("110-alerts.png",),
            ),
            Task(
                "Consulter les rapports généraux",
                "Analyser les effectifs, la présence et les indicateurs disponibles.",
                (
                    "Cliquez sur « Rapports ».",
                    "Choisissez la période et les filtres de parcours ou de classe.",
                    "Contrôlez les totaux et graphiques avant l’export.",
                    "Pour les états comptables détaillés, utilisez plutôt « Finance » puis « Rapports ».",
                ),
                ("111-reports.png",),
            ),
        ],
    ),
    (
        "12. Finance",
        "Encaisser, suivre les comptes élèves, gérer la trésorerie, la paie et la comptabilité.",
        [
            Task(
                "Lire le tableau de bord financier et réimprimer un reçu",
                "Consulter les encaissements et retrouver le reçu nominatif d’un paiement.",
                (
                    "Cliquez sur « Finance ».",
                    "Consultez les revenus, dépenses et solde de la période.",
                    "Dans l’historique, recherchez par reçu, élève, matricule, classe ou référence.",
                    "Cliquez sur l’icône de reçu de la ligne.",
                    "Utilisez « Télécharger PDF » ou « Imprimer » ; seul le reçu est produit, pas toute la page.",
                ),
                ("120-finance-overview.png", "120b-finance-payment-receipt.png"),
            ),
            Task(
                "Enregistrer un paiement",
                "Affecter immédiatement un versement à l’élève et au compte de trésorerie choisi.",
                (
                    "Dans « Finance », cliquez sur « Nouveau paiement ».",
                    "Recherchez l’élève et vérifiez son matricule et sa classe.",
                    "Renseignez le montant, la tranche ou le frais concerné, la date, le mode de paiement, la référence et le compte encaissé.",
                    "Validez, puis contrôlez le reçu et l’historique du compte élève.",
                ),
                ("121-finance-payment.png",),
                "Le compte choisi est impacté immédiatement ; vérifiez toujours Banque/Caisse avant de confirmer.",
            ),
            Task(
                "Consulter le compte d’un élève et produire un reçu consolidé",
                "Voir les frais, paiements et solde, même lorsque les versements ont été faits en plusieurs tranches.",
                (
                    "Dans « Finance », cliquez sur « Compte élève ».",
                    "Choisissez une classe pour afficher sa liste ; utilisez le nom ou le matricule pour affiner.",
                    "Cliquez sur l’élève et contrôlez Facturé, Payé, Solde dû, Crédit et chaque versement.",
                    "Cliquez sur « Préparer le reçu consolidé ».",
                    "Téléchargez le PDF ou imprimez le relevé regroupant tous les paiements.",
                ),
                ("122-finance-student-accounts.png", "122b-finance-student-history.png", "122c-finance-consolidated-receipt.png"),
            ),
            Task(
                "Gérer banques, caisse, dépôts, retraits et transferts",
                "Faire correspondre les soldes du système aux comptes bancaires et à la caisse.",
                (
                    "Dans « Finance », cliquez sur « Comptes & mouvements ».",
                    "Vérifiez les comptes Cash, BGFI Bank, Afriland, CCA, Regional et tout autre compte configuré.",
                    "Pour un dépôt ou retrait externe, choisissez l’opération, le compte, la contrepartie, le montant, le motif et la référence.",
                    "Pour un transfert, choisissez le compte source et le compte destination.",
                    "Enregistrez puis vérifiez les nouveaux soldes et la ligne immuable dans l’historique.",
                ),
                ("123-finance-treasury.png",),
            ),
            Task(
                "Configurer les types et plans de frais",
                "Définir ce qui est facturé, à quelles classes et selon quel échéancier.",
                (
                    "Ouvrez « Types de frais » et créez chaque frais avec son code, libellé et règles comptables.",
                    "Ouvrez « Plans de frais » et choisissez la session et le périmètre.",
                    "Ajoutez les montants, tranches, dates d’échéance et éventuelles dérogations.",
                    "Enregistrez le brouillon, contrôlez l’aperçu, puis publiez lorsque la configuration est exacte.",
                ),
                ("124-finance-fee-types.png", "125-finance-plans.png"),
            ),
            Task(
                "Générer les créances et contrôler les encaissements",
                "Créer les montants dus, suivre leur allocation et retrouver les documents financiers.",
                (
                    "Ouvrez « Créances » et lancez d’abord une prévisualisation pour la session et les classes choisies.",
                    "Contrôlez les élèves, montants et exceptions, puis générez les créances.",
                    "Ouvrez « Encaissements » pour vérifier l’allocation des paiements aux tranches et traiter les exceptions.",
                    "Ouvrez « Documents » pour retrouver factures et reçus générés.",
                ),
                ("126-finance-charges.png", "127-finance-collections.png", "128-finance-documents.png"),
            ),
            Task(
                "Préparer et payer le personnel",
                "Générer une paie, contrôler les montants et enregistrer le décaissement depuis un compte.",
                (
                    "Ouvrez « Paie du personnel ».",
                    "Créez la période de paie et chargez les employés concernés.",
                    "Contrôlez salaire, retenues et net à payer, puis validez la paie.",
                    "Choisissez le compte de trésorerie utilisé pour le paiement et enregistrez l’opération.",
                    "Générez et téléchargez les bulletins de paie.",
                ),
                ("129-finance-payroll.png",),
            ),
            Task(
                "Configurer le plan comptable et les règles de comptabilisation",
                "Préparer les comptes et les mappings utilisés par les opérations automatiques.",
                (
                    "Ouvrez « Comptabilité avancée ».",
                    "Dans « Comptes », créez ou contrôlez le plan comptable.",
                    "Dans « Mappings », associez chaque événement financier aux comptes débit et crédit attendus.",
                    "Vérifiez la checklist de préparation avant de poster des opérations.",
                ),
                ("130-finance-accounting.png", "131-finance-chart.png", "132-finance-mappings.png"),
            ),
            Task(
                "Gérer les périodes et journaux comptables",
                "Ouvrir les périodes, contrôler les écritures et clôturer au bon moment.",
                (
                    "Dans la comptabilité, ouvrez « Périodes » et générez les périodes liées à l’année scolaire.",
                    "Laissez la période ouverte pendant la saisie et clôturez-la seulement après rapprochement.",
                    "Ouvrez « Journaux » pour filtrer et consulter les écritures par date, source ou statut.",
                    "Ouvrez une écriture pour vérifier l’équilibre débit/crédit et sa pièce d’origine.",
                ),
                ("133-finance-periods.png", "134-finance-journals.png"),
            ),
            Task(
                "Contrôler la balance, le grand livre et le rapprochement",
                "Vérifier la cohérence comptable et résoudre les exceptions.",
                (
                    "Ouvrez « Balance » et contrôlez que les totaux débit et crédit sont égaux.",
                    "Ouvrez « Grand livre », choisissez un compte et vérifiez son détail chronologique.",
                    "Ouvrez « Rapprochement », comparez les opérations aux relevés et traitez chaque exception documentée.",
                    "Ne clôturez la période qu’après résolution ou justification des écarts.",
                ),
                ("135-finance-trial-balance.png", "136-finance-ledger.png", "137-finance-reconciliation.png"),
            ),
            Task(
                "Produire les rapports financiers",
                "Obtenir les synthèses officielles dans le bon contexte comptable.",
                (
                    "Dans « Finance », ouvrez « Rapports ».",
                    "Choisissez la période et le rapport souhaité.",
                    "Contrôlez les filtres, totaux et comptes inclus.",
                    "Exportez le document dans le format proposé et conservez le contexte de période avec le fichier.",
                ),
                ("138-finance-reports.png",),
            ),
        ],
    ),
    (
        "13. Espace parent",
        "Consulter les informations partagées par l’école et communiquer avec elle.",
        [
            Task(
                "Consulter l’enfant, son parcours et sa vie scolaire",
                "Voir les informations principales, les décisions officielles, la présence et les messages.",
                (
                    "Connectez-vous avec le compte parent et choisissez l’enfant en haut de page.",
                    "Utilisez « Vue d’ensemble » pour les indicateurs principaux.",
                    "Ouvrez « Parcours officiel » pour les décisions et résultats publiés.",
                    "Ouvrez « School life » pour la présence, la discipline et les communications partagées.",
                ),
                ("150-parent-home.png", "150b-parent-journey.png", "150c-parent-school-life.png"),
            ),
            Task(
                "Suivre les frais et les notes",
                "Consulter ce qui reste à payer et les résultats rendus visibles par l’école.",
                (
                    "Ouvrez « Frais & paiements » pour voir le montant facturé, déjà payé, le solde, les échéances et les moyens de paiement.",
                    "Téléchargez les factures ou reçus proposés.",
                    "Ouvrez « Notes » pour consulter les dernières notes et bulletins publiés.",
                    "Changez de programme lorsqu’un élève bilingue possède deux bulletins.",
                ),
                ("151-parent-fees.png", "152-parent-grades.png"),
            ),
            Task(
                "Consulter les fournitures et documents de l’école",
                "Retrouver les listes de classe et les documents partagés avec les familles.",
                (
                    "Ouvrez « Fournitures & manuels » et vérifiez les quantités, titres et indications facultatives.",
                    "Ouvrez « Documents de l’école » pour consulter ou télécharger les ressources publiées.",
                    "Choisissez le bon enfant si plusieurs enfants sont rattachés au compte.",
                ),
                ("153-parent-supplies.png", "153b-parent-library.png"),
            ),
            Task(
                "Envoyer une suggestion ou une question",
                "Transmettre un message structuré à l’école et suivre son traitement.",
                (
                    "Ouvrez « Boîte à suggestions ».",
                    "Choisissez la catégorie : suggestion, question, réclamation ou remerciement.",
                    "Rédigez un message précis, puis cliquez sur « Envoyer le message ».",
                    "Consultez « Mes messages » pour voir le statut de traitement.",
                ),
                ("154-parent-suggestions.png",),
            ),
        ],
    ),
]


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, color: str, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep(paragraph, next_paragraph: bool = False) -> None:
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = next_paragraph


def add_run_color(run, color: str, bold: bool = False, size: float | None = None) -> None:
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)


def add_callout(doc: Document, label: str, text: str, color: str = PALE_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, color)
    set_cell_border(cell, TEAL if color == PALE_TEAL else GOLD, "8")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}  ")
    add_run_color(r, TEAL if color == PALE_TEAL else NAVY, True, 9.5)
    r = p.add_run(text)
    add_run_color(r, INK, False, 9.5)


def add_screenshot(doc: Document, filename: str, caption: str) -> None:
    path = SCREENSHOTS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as im:
        width_px, height_px = im.size
    ratio = width_px / height_px
    # Tall workflow pages remain legible and fit inside A4; standard pages use
    # the full text width.  The original image is embedded, so zoom remains crisp.
    width_in = 6.85 if ratio >= 1.05 else min(6.4, 8.65 * ratio)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    keep(cap)
    run = cap.add_run(caption)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Guide complet BBC SMS"
    props.subject = "Guide utilisateur illustré de la plateforme scolaire"
    props.author = "Bayo Bilingual Complex"
    props.keywords = "BBC SMS, école, guide utilisateur, élèves, académique, finance"
    props.comments = "Validé sur l’application locale le 31 août 2026."


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 31, WHITE),
        ("Heading 1", 22, NAVY),
        ("Heading 2", 15, TEAL),
        ("Heading 3", 12, NAVY),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    styles["Heading 1"].paragraph_format.space_before = Pt(0)
    styles["Heading 1"].paragraph_format.space_after = Pt(9)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)
    styles["Caption"].font.name = "Aptos"
    styles["Caption"].font.size = Pt(8)
    styles["Caption"].font.italic = True

    if "Manual Step" not in styles:
        step = styles.add_style("Manual Step", WD_STYLE_TYPE.PARAGRAPH)
        step.base_style = styles["Normal"]
        step.font.name = "Aptos"
        step.font.size = Pt(10)
        step.paragraph_format.left_indent = Cm(0.55)
        step.paragraph_format.first_line_indent = Cm(-0.45)
        step.paragraph_format.space_after = Pt(4)


def setup_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(18.3))
    table.columns[0].width = Cm(12.5)
    table.columns[1].width = Cm(5.8)
    for cell in table.rows[0].cells:
        shade(cell, NAVY)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    add_run_color(left.add_run("BBC SMS  •  GUIDE UTILISATEUR"), WHITE, True, 8)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    add_run_color(right.add_run("31 AOÛT 2026"), GOLD, True, 8)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("Bayo Bilingual Complex  •  ")
    add_run_color(r, MUTED, False, 8)
    page_field(p)


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, NAVY)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("BBC SMS")
    add_run_color(r, GOLD, True, 17)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles["Title"]
    p.paragraph_format.space_after = Pt(10)
    p.add_run("Guide complet\nutilisateur")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Toutes les fonctions, expliquées pas à pas avec des captures réelles")
    add_run_color(r, WHITE, False, 12)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("VERSION FRANÇAISE  •  TESTÉE LOCALEMENT LE 31 AOÛT 2026")
    add_run_color(r, GOLD, True, 9)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(42)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Bayo Bilingual Complex")
    add_run_color(r, NAVY, True, 14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Guide destiné aux utilisateurs de la plateforme. Les fonctions d’administration des rôles, des accès et de la messagerie technique sont volontairement exclues.")
    add_run_color(r, MUTED, False, 9.5)
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    h = doc.add_heading("Comment utiliser ce guide", level=1)
    keep(h, True)
    doc.add_paragraph(
        "Chaque fiche répond à une question simple : que voulez-vous faire, où cliquer et quoi vérifier avant de confirmer. "
        "Les captures proviennent de l’application locale réellement testée ; les noms et montants sont des données de démonstration."
    )
    add_callout(
        doc,
        "RÈGLE SIMPLE",
        "Commencez par Applications, ouvrez le module, choisissez le bon parcours ou la bonne classe, puis vérifiez le résumé avant d’enregistrer.",
    )

    doc.add_heading("Repères visuels", level=2)
    rows = [
        ("Bouton foncé ou coloré", "Action principale à exécuter."),
        ("Liste déroulante", "Choix du parcours, de la classe, de la période, de la matière ou du compte."),
        ("Badge vert", "État terminé, payé, publié ou validé."),
        ("Badge orange / rouge", "Élément incomplet, en retard ou à contrôler."),
        ("Bouton grisé", "Action indisponible dans l’état actuel ; contrôlez les prérequis indiqués à l’écran."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(5.2)
    table.columns[1].width = Cm(12.8)
    set_repeat_table_header(table.rows[0])
    for i, title in enumerate(("À l’écran", "Signification")):
        cell = table.cell(0, i)
        shade(cell, NAVY)
        add_run_color(cell.paragraphs[0].add_run(title), WHITE, True, 9)
    for a, b in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        shade(cells[0], PALE)

    doc.add_heading("Sommaire", level=2)
    for title, subtitle, _ in CHAPTERS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title)
        add_run_color(r, NAVY, True, 10.5)
        r = p.add_run(f" — {subtitle}")
        add_run_color(r, MUTED, False, 9.5)
    doc.add_page_break()


def add_task(doc: Document, task: Task, manifest_titles: dict[str, str]) -> None:
    heading = doc.add_heading(task.title, level=2)
    keep(heading, True)
    add_callout(doc, "OBJECTIF", task.goal)
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(6)
    label.paragraph_format.space_after = Pt(3)
    keep(label, True)
    add_run_color(label.add_run("Étapes"), NAVY, True, 10.5)
    for index, step in enumerate(task.steps, 1):
        p = doc.add_paragraph(style="Manual Step")
        p.add_run(f"{index}. ").bold = True
        p.add_run(step)
    if task.tip:
        add_callout(doc, "À RETENIR", task.tip, "FFF5D9")
    for image_name in task.images:
        caption = manifest_titles.get(image_name, image_name)
        add_screenshot(doc, image_name, caption)


def build() -> None:
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    if manifest.get("failed") != 0 or manifest.get("passed") != manifest.get("total"):
        raise RuntimeError("The browser manifest is not fully green; refusing to build the final manual.")
    manifest_titles = {row["file"]: row["title"] for row in manifest["captures"] if row.get("file")}
    required = {image for _, _, tasks in CHAPTERS for task in tasks for image in task.images}
    missing = sorted(required - set(manifest_titles))
    if missing:
        raise RuntimeError(f"Screenshots absent from the green manifest: {missing}")

    doc = Document()
    set_core_properties(doc)
    setup_styles(doc)
    setup_sections(doc)
    add_cover(doc)
    add_front_matter(doc)

    for chapter_title, subtitle, tasks in CHAPTERS:
        h = doc.add_heading(chapter_title, level=1)
        keep(h, True)
        p = doc.add_paragraph(subtitle)
        p.paragraph_format.space_after = Pt(8)
        p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        p.runs[0].italic = True
        for task in tasks:
            add_task(doc, task, manifest_titles)
        doc.add_page_break()

    doc.add_heading("Checklist avant de confirmer une opération", level=1)
    checklist = [
        "Le bon parcours, la bonne classe, la bonne période et le bon élève sont sélectionnés.",
        "Les totaux affichés correspondent au résultat attendu.",
        "Les pièces jointes s’ouvrent correctement.",
        "Les montants financiers utilisent le bon compte Banque/Caisse et une référence traçable.",
        "Les brouillons ont été relus avant validation, publication, finalisation ou clôture.",
        "Le document exporté ou imprimé a été ouvert une fois pour contrôler sa mise en page.",
    ]
    for item in checklist:
        p = doc.add_paragraph(style="Manual Step")
        add_run_color(p.add_run("☐  "), TEAL, True, 12)
        p.add_run(item)
    add_callout(
        doc,
        "SUPPORT",
        "En cas de blocage, notez le module, la classe, la période, le nom de l’élève ou de l’employé et le message affiché. Une capture de la page entière accélère le diagnostic.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(json.dumps({
        "output": str(OUT),
        "chapters": len(CHAPTERS),
        "tasks": sum(len(tasks) for _, _, tasks in CHAPTERS),
        "screenshots": len(required),
        "manifestTotal": manifest["total"],
    }, ensure_ascii=False))


if __name__ == "__main__":
    build()
